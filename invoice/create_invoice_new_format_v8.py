#!/usr/bin/env python
# -*- coding: utf8 -*-


from pdfminer.layout import LAParams, LTTextBox
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfpage import PDFPage
from pdfminer.layout import LTTextBoxHorizontal
import os
from os import listdir, pardir
from os.path import isfile, join, abspath
from docx.api import Document
from docx.shared import Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime
import sys
from easygui import fileopenbox
import pandas as pd
from openpyxl import load_workbook
from copy import copy
from openpyxl.utils import get_column_letter
import os.path
import re
import unicodedata


def get_contents(filename):
  # found this online
  # outputs a single list of strings from a pdf
  rsrcmgr = PDFResourceManager()
  laparams = LAParams()
  device = PDFPageAggregator(rsrcmgr, laparams=laparams)
  interpreter = PDFPageInterpreter(rsrcmgr, device)
  pdf_file_instance = open(filename, 'rb')
  total_text = []
  
  # Boxes to look for info
  #xmin,xmax, ymin,ymax
  xy = [
    (325,335, 750,805), # commande_n
    (190,210, 550,564), # reference_n
    (31,32, 300,506), # tasks
    (500,515, 300,506), # prices
    (500,550, 145,216) # total ht, total ht net, total tva, net a payer
   ]
  #xy = [
  #  (325,335, 800,805), # commande_n
  #  (190,210, 560,564), # reference_n
  #  (31,32, 300,506), # tasks
  #  (509,515, 300,506), # prices
  #  (523,535, 145,216) # total ht, total ht net, total tva, net a payer
  # ]
   
  res = [' ']*len(xy)

  for page in PDFPage.get_pages(pdf_file_instance, maxpages=1):
      interpreter.process_page(page)
      layout = device.get_result()
      for lobj in layout:
        if isinstance(lobj, LTTextBox):
          x, y, text = lobj.bbox[0], lobj.bbox[3], lobj.get_text()
          #print('At %r is text: %s' % ((x, y), text))  
          if any(a <= x <= b and c <=y <=d for a,b,c,d in xy):
                    
            for idx, (a,b,c,d) in enumerate(xy):
              if a <= x <= b and c <=y <=d:
                res[idx] = '\n'.join([res[idx],text]).strip()
              #print('%r text: %s' % ((x, y), text))
                
            #print('At %r is text: %s' % ((x, y), text))
  pdf_file_instance.close()
           
  try:
    commande_n = res[0].split('\n')[-1]
  except:
    commande_n = ''
  try:
    reference_n = res[1]
  except:
    reference_n = ''
  try:
    tasks = res[2]
  except:
    tasks = ''
  try:
    prices = res[3]
  except:
    prices = ''
  try:
    total_ht,_,total_tva,total_ttc = res[4].split('\n')
    #note need to deal with total_ht.replace(u'\xa0', u'').
    #this should be done in the update_facture function

  except:
    total_ht,total_tva,total_ttc = '','',''

  
  return(commande_n, reference_n, total_ht, total_tva, total_ttc, tasks, prices)


# function to generate the facture number
def create_facture_n():

  mypath = os.getcwd()

  #find all docx files
  onlyfiles = [f for f in listdir(mypath) if (isfile(join(mypath, f)) and f.endswith(".docx"))]

  # find the existing facture numbers
  # remove any numbers that dont make sense
  # i.e. numbers should be in the range yymm00
  # or max 6 digits
  a1 = []
  for i in onlyfiles:
    t = [int(s) for s in i.split("-") if s.isdigit() if int(s)<1000000 if int(s)>=99999]
    a1.append(t)
  # remove empty lists within the list
  a = [x for x in a1 if x]

  facture_n=0

  # if previous facture number exists, increment by one
  # otherwise create a new number
  if (len(a)>=1):
    facture_n = max(a)[0]
    # increment by 1 from the maximum
    facture_n += 1
  else:
    temp = datetime.datetime.today().strftime('%y%m')
    facture_n = temp + "01"
    facture_n = int(facture_n)

  text = str(facture_n) + '-facturation-VIART-Alto.docx'
  print("   Invoice number %d " %facture_n)
  return(facture_n, text)

def update_facture(facture_n, reference_n, HT, frais=0):

  # function to convert row,col to A3, B4, etc
  def xlref(row, column, zero_indexed=True):
    if zero_indexed:
        row += 1
        column += 1
    return get_column_letter(column) + str(row)



  # open file, which is two directories below
  infile = 'Factures indépendant.xlsx'
  filepath = os.path.abspath(os.path.join(os.getcwd(),"../../.."))
  file = os.path.join(filepath, infile)
  #print(file)

  xl = load_workbook(file)


  # find first empty row after row 10
  # since there are a few empty rows at the beginning...
  ws = xl.worksheets[0]
  cell_list = []
  for cell in ws["C"]:
    if cell.value is None:
        cell_list.append(cell.row)

  cell_list2 = [i for i in cell_list if i>=10]
  first_row = cell_list2[0]


  # values to add to the excel file
  #N° facture   Client  HT  Frais   Total HT    TVA TTC Mission

  #Replicate the formulas in the worksheets
  # OR use regular expression to increment all numbers by 1
  #import re
  #re.sub('\d+',lambda x:str(int(x.group())+1),test_text)

  #totalHT=SUM(C35:D35)
  totht = "=SUM(" + xlref(first_row, 3, False) + ":" + xlref(first_row, 4, False) + ")"
  #TVA=E35/5
  tva = "=" + xlref(first_row, 5, False) + "/5"
  #TTC=SUM(E35:F35)
  ttc = "=SUM(" + xlref(first_row, 5, False) + ":" + xlref(first_row, 6, False) + ")"

  
  # save the new line data into a dataframe
  newdf = (facture_n, "ALTO", float(HT.replace(u'\xa0', u'').replace(' ','').replace(',','.')), frais, totht, tva, ttc, reference_n)


  # add new row cell by cell and copy formatting
  for i,value in enumerate(newdf):
    ws.cell(column=i+1, row=first_row).value=value
    ws.cell(column=i+1, row=first_row)._style = copy(ws.cell(column=i+1, row=(first_row-2)))._style

  #before saving, back up old file to new folder
  backup_f = 'Factures_backup'
  os.path.abspath(os.path.join(os.getcwd(),"../../.."))
  backup = os.path.join(filepath, backup_f)

  if not os.path.isdir(backup):
    os.makedirs(backup)
    #print("Home directory %s was created." %home_dir)


  from shutil import copyfile

  temp = datetime.datetime.today().strftime('%y%m%d')
  backup_name = 'Factures indépendant ' + temp + '.xlsx'
  copyfile(file, os.path.join(backup, backup_name))

  xl.save(file)





# function to move the pdf to current directory
def move_pdf(infile, reference_n):
  from shutil import copyfile


  #start = infile.find('\Com')
  #end = infile.find('.pdf')
  #temp_name = infile[start+1:end] + ' '
  temp_name = infile.split('\\')[-1].split('.pdf')[0]
  
  #replaced = re.sub('\([0-9]\)', '', temp_name)
    
  new_name = temp_name  + str(reference_n) + '.pdf'
  new_file_path = os.path.abspath(os.path.join(os.getcwd(), new_name))

  copyfile(infile, new_file_path)
  #Deleted at the end of program instead
  #print("   Deleting file: ",infile)
  #os.remove(infile)
  

  open_file_for_tom(r'C:\Program Files (x86)\Adobe\Acrobat Reader DC\Reader\AcroRd32.exe', new_file_path)


def open_file_for_tom(program, fn):
  import subprocess

  try:
    subprocess.Popen([program, fn])
  except:
    print(program, fn)
    print("Failed for ", fn)


#function to generate the document
def make_docx(infile):

  #generate facture number and filename for saving
  facture_n, outfile = create_facture_n()

  #step 1: get info from text
  commande_n, reference_n, total_ht, total_tva, total_ttc, tasks, prices = get_contents(infile)
  

  #step 2b: append the reference number to the end of the pdf file
  move_pdf(infile, reference_n)

  #step 3: make the doc file
  #prepare shadings for 2 final table cells
  shading_1 = parse_xml(r'<w:shd {} w:fill="#99ccff"/>'.format(nsdecls('w')))
  shading_2 = parse_xml(r'<w:shd {} w:fill="#99ccff"/>'.format(nsdecls('w')))

  newdoc = Document()

  #Title info
  p1 = newdoc.add_paragraph()
  p1.add_run('Tom Viart').bold = True
  p1.add_run('''
132 rue du Chemin Vert
75011 PARIS
Tél : 06 47 59 86 17''')
  p1.add_run('''
Identifiant SIRET : 75105892600021
N° TVA : FR 68 751058926''').bold = True

  p2 = newdoc.add_paragraph('À l’attention d’')
  p2.add_run('ALTO').bold = True
  p2.add_run('''
99, rue du Faubourg Saint Martin
75010 PARIS
  ''')
  p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

  today = datetime.datetime.today().strftime('%d/%m/%Y')
  p3 = newdoc.add_paragraph()
  p3.add_run('FACTURE %d' % facture_n).bold = True
  p3.add_run('''
  En date du %s''' % today)
  p3.alignment = WD_ALIGN_PARAGRAPH.CENTER


  #make the table of prices
  table = newdoc.add_table(rows=5, cols=2)
  table.style = 'Table Grid'
  table.alignment = WD_ALIGN_PARAGRAPH.CENTER


  # format cell widths
  for cell in table.columns[0].cells:
    cell.width = Inches(4.0)
  for cell in table.columns[1].cells:
    cell.width = Inches(0.9)

  a = table.cell(0, 0)
  b = table.cell(0,1)
  A = a.merge(b)

  tt = 'Commande N° %s \nRéférence %s\n' % (commande_n,reference_n)
  A2 = A.add_paragraph(tt)
  A2.alignment = WD_ALIGN_PARAGRAPH.CENTER

  c = table.cell(1,0)
  c.text = tasks
  d = table.cell(1,1)
  d.text = prices
  d.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  e = table.cell(2,0)
  e.text = "Total HT"
  e.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  e2 = table.cell(2,1)
  e2.text = total_ht
  e2.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  f = table.cell(3,0)
  f.text = "TVA 20 %"
  f.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  f2 = table.cell(3,1)
  f2.text = total_tva
  f2.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  g = table.cell(4,0)
  g.text = "Total TTC"
  g.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  g._tc.get_or_add_tcPr().append(shading_1)
  g2 = table.cell(4,1)
  g2.text = total_ttc
  g2.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  g2._tc.get_or_add_tcPr().append(shading_2)




  ### End of document -- does not change
  newdoc.add_paragraph('''\n\nValeur en votre aimable règlement à réception (RIB ci-joint)
  ''')
  newdoc.add_paragraph('''\nRELEVE D’IDENTITE BANCAIRE''')

  table2 = newdoc.add_table(rows=2, cols=4)
  table2.style = 'Table Grid'

  a = table2.cell(0,0)
  a.text = "Banque"
  a = table2.cell(0,1)
  a.text = "Guichet"
  a = table2.cell(0,2)
  a.text = "Numéro de compte"
  a = table2.cell(0,3)
  a.text = "Clé RIB"

  b = table2.cell(1,0)
  b.text = "10278"
  b = table2.cell(1,1)
  b.text = "06041"
  b = table2.cell(1,2)
  b.text = "00020864601"
  b = table2.cell(1,3)
  b.text = "54"


  newdoc.add_paragraph('''
  IDENTIFICATION INTERNATIONALE''')

  table3 = newdoc.add_table(rows=1, cols=1)
  table3.style = 'Table Grid'
  a = table3.cell(0,0)
  a.text = '''IBAN : FR76 1027 8060 4100 0208 6460 154
BIC : CMCIFR2A'''


  newdoc.save(outfile)
  out_location = os.path.abspath(os.path.join(os.getcwd(), outfile))

  # Update the facture xls document using update facture() function
  try:
    update_facture(facture_n, reference_n, total_ht)
    print("   Created file: ",outfile)
  except:
    print("!! Error making excel file for {}".format(infile))
    #print("\n".join([commande_n, reference_n, total_ht, total_tva, total_ttc, tasks, prices]))
    

  #open DOC file
  open_file_for_tom(r'C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE', out_location)
  
  #create a warning message if not all values were found
  if '' in [commande_n, reference_n, total_ht, total_tva, total_ttc, tasks, prices]:
    from easygui import msgbox
    msgbox(msg='Error: program did not correctly find all values', title='Error', ok_button='OK', image=None, root=None)
    errors = pd.DataFrame([commande_n, reference_n, total_ht, total_tva, total_ttc, tasks, prices],columns = ['commande_n', 'reference_n', 'total_ht', 'total_tva', 'total_ttc', 'tasks', 'prices'])
    errors.to_csv('errors.csv',index=False)
    




# main function
def main():
    # command line arguments
    i=1
    openfiles = fileopenbox("Choose one or more ALTO pdf files","Invoice Generation",default=r"C:\Users\tomvi\Desktop\*.pdf",filetypes="*.pdf",multiple=True)
    for arg in openfiles:
      print("Generating Document %d..." %i)
      make_docx(arg)      
      i+=1
    #open excel file at the end
    open_file_for_tom(r'C:\Program Files\Microsoft Office\root\Office16\EXCEL.EXE', r'C:\Users\tomvi\Desktop\Factures\Factures indépendant.xlsx')
    print()
    print("Deleting: ")
    for arg in openfiles:
      print("   ",arg)
      os.remove(arg)
	
if __name__ == "__main__":
    main()
